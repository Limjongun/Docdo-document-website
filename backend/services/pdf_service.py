import fitz  # PyMuPDF
import io
from pathlib import Path
from PIL import Image
import uuid
from docx import Document
from docx.shared import Inches
import PyPDF2


def pdf_to_word(pdf_bytes: bytes) -> bytes:
    """Convert PDF pages to a .docx file (text extraction per page)."""
    doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    word_doc = Document()
    word_doc.add_heading("Converted Document", level=1)

    for page_num, page in enumerate(doc_pdf, start=1):
        word_doc.add_heading(f"Page {page_num}", level=2)
        text = page.get_text("text")
        word_doc.add_paragraph(text if text.strip() else "[No text found on this page]")
        word_doc.add_paragraph()  # spacer

    buf = io.BytesIO()
    word_doc.save(buf)
    return buf.getvalue()


def add_watermark(pdf_bytes: bytes, watermark_text: str) -> bytes:
    """Add a semi-transparent text watermark to every PDF page."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        page.insert_text(
            point=(page.rect.width / 4, page.rect.height / 2),
            text=watermark_text,
            fontsize=48,
            color=(0.8, 0.8, 0.8),
            rotate=45,
        )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def set_pdf_password(pdf_bytes: bytes, password: str) -> bytes:
    """Encrypt a PDF with a user password."""
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    writer = PyPDF2.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def remove_pdf_password(pdf_bytes: bytes, password: str) -> bytes:
    """Remove password from an encrypted PDF."""
    reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    if reader.is_encrypted:
        reader.decrypt(password)
    writer = PyPDF2.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def rotate_pdf(pdf_bytes: bytes, degrees: int) -> bytes:
    """Rotate all pages of a PDF by given degrees (90, 180, 270)."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        page.set_rotation(degrees)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def pdf_to_images(pdf_bytes: bytes, fmt: str = "png") -> list[bytes]:
    """Convert each PDF page to an image and return list of image bytes."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    for page in doc:
        mat = fitz.Matrix(2.0, 2.0)  # 2x scale = 144 DPI
        pix = page.get_pixmap(matrix=mat)
        images.append(pix.tobytes(fmt))
    return images


def embed_signature(pdf_bytes: bytes, sig_bytes: bytes, page_num: int = 0) -> bytes:
    """Embed a signature image onto a specified PDF page."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num]
    # Place signature at bottom-right
    rect = fitz.Rect(
        page.rect.width - 200,
        page.rect.height - 100,
        page.rect.width - 20,
        page.rect.height - 20,
    )
    page.insert_image(rect, stream=sig_bytes)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
